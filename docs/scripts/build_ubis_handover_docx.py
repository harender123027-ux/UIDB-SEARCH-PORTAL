#!/usr/bin/env python3
"""
Build a single Word handover package from UBIS docs (for police / management).
Requires: python-docx (pip install python-docx)

Output (under docs/):
  UBIS_Handover_Package.docx

Optional: run `textutil -convert doc UBIS_Handover_Package.docx` on macOS for .doc
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

DOCS = Path(__file__).resolve().parents[1]
OUT_DOCX = DOCS / "UBIS_Handover_Package.docx"
GURUGRAM_DIR = DOCS / "HANDOVER_GURUGRAM"
OUT_DOCX_GURUGRAM = DOCS / "UBIS_Gurugram_Handover_Package.docx"


def add_formatted_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    """Inline **bold** and strip simple `code` backticks; drop markdown links to label only."""
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    remaining = text.strip()
    while remaining:
        m = re.search(r"\*\*([^*]+)\*\*", remaining)
        if not m:
            chunk = remaining.replace("`", "")
            if chunk:
                p.add_run(chunk)
            break
        before = remaining[: m.start()].replace("`", "")
        if before:
            p.add_run(before)
        r = p.add_run(m.group(1))
        r.bold = True
        remaining = remaining[m.end() :]


def is_table_separator(line: str) -> bool:
    s = line.strip()
    if not s.startswith("|"):
        return False
    core = s.strip("|").replace(" ", "")
    return bool(core) and all(c in "-:|" for c in core)


def flush_table(doc: Document, rows: list[list[str]]) -> None:
    if len(rows) < 2:
        for r in rows:
            add_formatted_paragraph(doc, " | ".join(r))
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = row[j] if j < len(row) else ""
            table.rows[i].cells[j].text = re.sub(r"\*\*([^*]+)\*\*", r"\1", cell)
            for p in table.rows[i].cells[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)


def append_markdown_file(doc: Document, path: Path) -> None:
    text = path.read_text(encoding="utf-8")
    lines = text.splitlines()
    table_buf: list[list[str]] = []

    def flush_buf() -> None:
        nonlocal table_buf
        if table_buf:
            flush_table(doc, table_buf)
            table_buf = []

    for line in lines:
        if line.strip() == "---":
            flush_buf()
            continue
        stripped = line.strip()
        if stripped.startswith("|") and not is_table_separator(stripped):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            table_buf.append(cells)
            continue
        flush_buf()

        if not stripped:
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=0)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=3)
        else:
            add_formatted_paragraph(doc, stripped)

    flush_buf()


def append_management_brief_txt(doc: Document, path: Path) -> None:
    """Formal ASCII brief: numbered sections, metadata lines, bullets."""
    raw = path.read_text(encoding="utf-8")
    seen_main_title = False
    for line in raw.splitlines():
        s = line.strip()
        if not s:
            continue
        if s == "---":
            doc.add_paragraph(" ")
            continue
        if s.startswith("=" * 10):
            continue
        if not seen_main_title and s.startswith("UBIS —") and "BRIEF" in s.upper():
            doc.add_heading(s, level=0)
            seen_main_title = True
            continue
        if seen_main_title and s.startswith("Unidentified Body Identification System"):
            doc.add_paragraph(s)
            continue
        if re.match(r"^(Reference|Date|Distribution|Prepared by):", s):
            doc.add_paragraph(s)
            continue
        if re.match(r"^\d+\.\s+", s):
            doc.add_heading(s, level=1)
            continue
        if s.startswith("- "):
            add_formatted_paragraph(doc, "• " + s[2:].strip())
        elif "end of document" in s.lower():
            doc.add_paragraph(s)
        else:
            add_formatted_paragraph(doc, s)


def build() -> Path:
    doc = Document()
    sect = doc.sections[0]
    sect.top_margin = Inches(0.75)
    sect.bottom_margin = Inches(0.75)

    title = doc.add_heading("UBIS — Handover documentation package", 0)
    title.runs[0].font.size = Pt(22)
    doc.add_paragraph(
        "Compiled for police and senior stakeholders. This volume combines the executive brief "
        "and UAT / sign-off material. Regenerate from source files after updates."
    )

    doc.add_heading("Part A — Executive status and handover brief", level=1)
    append_management_brief_txt(doc, DOCS / "SENIOR_MANAGEMENT_PROGRESS_NOTE.txt")

    doc.add_page_break()
    doc.add_heading("Part B — UAT, testing, police sign-off & cloud sizing", level=1)
    append_markdown_file(doc, DOCS / "UAT_AND_POLICE_SIGNOFF.md")

    doc.add_page_break()
    doc.add_heading("Part C — Reference (titles only)", level=1)
    doc.add_paragraph(
        "Full technical detail remains in the repository: API_REFERENCE.md, POLICE_STATION_BULK_DATA_GUIDE.md, "
        "SYSTEM_DESIGN.md, DATA_INTERACTIONS.md, TESTING_GUIDE.md, docs/HANDOVER_GURUGRAM/README.md."
    )

    doc.save(OUT_DOCX)
    return OUT_DOCX


def build_gurugram() -> Path:
    """Compile the on-prem IT/operations + training book for Gurugram Police IT."""
    if not GURUGRAM_DIR.is_dir():
        raise FileNotFoundError(f"Gurugram handover dir missing: {GURUGRAM_DIR}")

    doc = Document()
    sect = doc.sections[0]
    sect.top_margin = Inches(0.75)
    sect.bottom_margin = Inches(0.75)

    title = doc.add_heading("UBIS — Gurugram Police on-prem handover book", 0)
    title.runs[0].font.size = Pt(22)
    doc.add_paragraph(
        "Volume II of the UBIS handover. This book is for the Gurugram Police IT team and trainers. "
        "It covers everything needed to install, operate, secure, train, and accept UBIS on a single "
        "police-controlled Linux server, with no cloud dependency. Volume I (the executive / UAT book, "
        "UBIS_Handover_Package.docx) covers programme governance and sign-off."
    )
    doc.add_paragraph(
        "Each section below is a self-contained chapter; cross-references use the chapter numbers "
        "(e.g. 'see Chapter 05'). The same files exist as plain Markdown under docs/HANDOVER_GURUGRAM/ "
        "in the source tree."
    )

    paths = sorted(p for p in GURUGRAM_DIR.glob("*.md") if not p.name.startswith("."))
    readme_paths = [p for p in paths if p.name == "README.md"]
    other_paths = [p for p in paths if p.name != "README.md"]
    chapters = readme_paths + other_paths
    for chapter_path in chapters:
        doc.add_page_break()
        append_markdown_file(doc, chapter_path)

    doc.save(OUT_DOCX_GURUGRAM)
    return OUT_DOCX_GURUGRAM


if __name__ == "__main__":
    paths = []
    try:
        paths.append(build())
    except Exception as e:
        print("Volume I build failed:", e, file=sys.stderr)
        sys.exit(1)
    try:
        paths.append(build_gurugram())
    except Exception as e:
        print("Volume II build failed:", e, file=sys.stderr)
        sys.exit(1)
    for p in paths:
        print("Wrote:", p)
